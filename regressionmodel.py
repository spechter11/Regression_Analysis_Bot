import numpy as np
import pandas as pd
import logging
import os
import re
from docx import Document
from typing import List, Dict, Optional, Tuple, Union, Any
from datetime import datetime
from scipy import stats
from sklearn.linear_model import LinearRegression, Ridge, Lasso
from sklearn.model_selection import train_test_split
from sklearn.compose import ColumnTransformer, make_column_selector
from sklearn.preprocessing import OneHotEncoder
from sklearn.metrics import mean_squared_error
from sklearn.preprocessing import StandardScaler
import matplotlib.pyplot as plt
import seaborn as sns
import statsmodels.api as sm

logging.basicConfig(level=logging.INFO)

class DataHandler:
    def load_data(self, file_path: str, required_columns: Optional[List[str]] = None, sheet_name: str = 'Data') -> pd.DataFrame:
        try:
            if file_path.endswith('.xlsx'):
                data = pd.read_excel(file_path, sheet_name=sheet_name)
            elif file_path.endswith('.csv'):
                data = pd.read_csv(file_path)
            else:
                logging.error(f"Unsupported file type: {file_path}")
                raise ValueError(f"Unsupported file type: {file_path}")

            # Convert column names to string
            data.columns = data.columns.astype(str)

            if required_columns:
                self._validate_columns(data, required_columns)

            return data

        except Exception as e:
            logging.error(f"An error occurred while loading data: {str(e)}")
            raise

    def _validate_columns(self, data: pd.DataFrame, required_columns: List[str]) -> None:
        missing_columns = [col for col in required_columns if col not in data.columns]
        if missing_columns:
            logging.error(f"The following columns are missing from the DataFrame: {missing_columns}")
            raise KeyError(f"The following columns are missing: {missing_columns}")

    def preprocess_data(self, X: pd.DataFrame, y: pd.Series) -> Tuple[pd.DataFrame, pd.DataFrame, pd.Series, pd.Series, List[str]]:
        try:
            updated_feature_names = X.columns.tolist()

            # Check if 'season' column exists for one-hot encoding
            if 'season' in X.columns:
                # One-hot encoding setup
                column_transformer = ColumnTransformer(
                    [('season_encoder', OneHotEncoder(), ['season'])],
                    remainder='passthrough'
                )
                # Perform one-hot encoding
                X_encoded = column_transformer.fit_transform(X)
                # Retrieve the one-hot encoded feature names
                one_hot_columns = column_transformer.named_transformers_['season_encoder'].get_feature_names_out(input_features=['season'])
                updated_feature_names = list(one_hot_columns) + updated_feature_names
            else:
                X_encoded = X.values

            return pd.DataFrame(X_encoded, columns=updated_feature_names), y, updated_feature_names
        except Exception as e:
            logging.error(f"An error occurred during preprocessing: {str(e)}")
            raise



    def detect_outliers(self, X: pd.DataFrame) -> Union[List[int], None]:
        """
        Detects outliers in the dataset based on the Z-score.

        :param X: Independent variables.
        :return: Indices of the detected outliers.
        """
        try:
            z_score = np.abs(stats.zscore(X))
            outliers = np.where(z_score > 2)
            return outliers[0].tolist()
        except Exception as e:
            logging.error(f"An error occurred while detecting outliers: {str(e)}")
            raise

class ReportGenerator:
    @staticmethod
    def save_to_word_doc(text: str, filename_prefix: str = 'regressionanalysis') -> None:
        """
        Saves the provided text to a Word document.

        :param text: Text to be saved in the document.
        :param filename_prefix: Prefix for the filename.
        """
        try:
            # Create a new Document
            doc = Document()
            
            # Add a Title
            doc.add_heading('Regression Analysis Report', 0)
            
            # Add the insights
            doc.add_paragraph(text)
            
            # Generate the filename
            timestamp = datetime.now().strftime('%m%d%y')
            filename = f"{filename_prefix}{timestamp}.docx"
            
            # Check for filename collision
            counter = 1
            while os.path.exists(filename):
                filename = f"{filename_prefix}{timestamp}_{counter}.docx"
                counter += 1
            
            # Save the document
            doc.save(filename)
            logging.info(f"Report saved as {filename}")
        
        except Exception as e:
            logging.error(f"Could not save the report: {str(e)}")

class StatisticalAnalysis:
    def __init__(self, X: pd.DataFrame, y: pd.Series) -> None:
        """
        Initialize the StatisticalAnalysis class with independent variables (X) and dependent variable (y).

        :param X: Independent variables (DataFrame).
        :param y: Dependent variable (Series).
        """
        self.X = X
        self.y = y

    def perform_analysis(self, feature_names: List[str]) -> sm.regression.linear_model.RegressionResultsWrapper:
        """
        Perform regression analysis and log the summary statistics.

        :param feature_names: List of feature names after preprocessing.
        :return: Fitted OLS model object.
        """
        try:
            # Adding a constant for the intercept term
            X_with_const = sm.add_constant(self.X)

            # Set the correct column names, including the intercept
            X_with_const.columns = ['const'] + feature_names

            # Fitting the Ordinary Least Squares (OLS) model
            model = sm.OLS(self.y, X_with_const).fit()

            # Logging a summary of the regression results
            logging.info(model.summary())

            return model
        
        except Exception as e:
            logging.error(f"Failed to perform analysis: {str(e)}")
            return None


class BaseInterpreter:
    def __init__(self, regression_output: str) -> None:
        self.output = regression_output
        self.interpretation = ""

    def parse(self, pattern: str, all_groups: bool = False) -> str:
        """
        Extracts data using a regex pattern from the regression output.
        
        :param pattern: Regular expression pattern to search for.
        :param all_groups: If True, returns all captured groups. Otherwise, returns just the first group.
        :return: Captured group(s) from the regex search.
        """
        try:
            match = re.search(pattern, self.output)
            if not match:
                raise ValueError(f"Pattern '{pattern}' not found in the regression output.")
            
            return match.groups() if all_groups else match.group(1)
        except Exception as e:
            logging.error(f"Failed to parse pattern: {str(e)}")

    def interpret(self) -> None:
        """
        Interpretation method that should be implemented by derived classes.
        """
        raise NotImplementedError("Interpret method not implemented in the derived class.")

    def get_interpretation(self) -> str:
        """
        Returns the constructed interpretation.
        
        :return: Interpretation string.
        """
        return self.interpretation

class ModelFitInterpreter(BaseInterpreter):
    def interpret(self) -> str:
        """
        Combines the interpretations of R-squared, Adjusted R-squared, and F-statistic.
        
        :return: Complete interpretation string.
        """
        try:
            self.interpretation += self.interpret_r2()
            self.interpretation += self.interpret_adj_r2()
            self.interpretation += self.interpret_f_statistic()
            return self.interpretation
        except Exception as e:
            logging.error(f"Failed to generate interpretation: {str(e)}")
            return ""

    def interpret_r2(self) -> str:
        """
        Interprets the R-squared value.
        
        :return: Interpretation of R-squared.
        """
        r2 = float(self.parse(r"R-squared:\s+(\d+\.\d+)"))
        return f"R-squared interpretation based on value: {r2}.\n"

    def interpret_adj_r2(self) -> str:
        """
        Interprets the Adjusted R-squared value.
        
        :return: Interpretation of Adjusted R-squared.
        """
        adj_r2 = float(self.parse(r"Adj. R-squared:\s+(\d+\.\d+)"))
        return f"Adjusted R-squared interpretation based on value: {adj_r2}.\n"

    def interpret_f_statistic(self) -> str:
        """
        Interprets the F-statistic p-value.
        
        :return: Interpretation of F-statistic p-value.
        """
        f_stat_p_value = float(self.parse(r"Prob \(F-statistic\):\s+(\d+\.\d+)"))
        return f"F-statistic p-value interpretation: {f_stat_p_value}.\n"


# Variable Interpreter
class VariableInterpreter(BaseInterpreter):
    def __init__(self, output: str) -> None:
        super().__init__(output)
    
    def interpret(self, feature_names: List[str]) -> str:
        """
        Interprets the coefficients of the variables.

        :param feature_names: List of feature names.
        :return: Complete interpretation string.
        """
        try:
            coef_data = self.parse_coef_data()
            for coef in coef_data:
                self.interpretation += self.interpret_coefficient(coef, feature_names)
            return self.interpretation
        except Exception as e:
            logging.error(f"Failed to generate interpretation: {str(e)}")
            return ""


    def parse_coef_data(self) -> List[Dict[str, float]]:
        """
        Extracts coefficient data from the regression output.
        
        :return: List of dictionaries containing coefficient data.
        """
        pattern = r"(\w+)\s+([\d.-]+)\s+([\d.-]+)\s+([\d.-]+)\s+([\d.-]+)"
        matches = re.findall(pattern, self.output)
        return [{'variable': match[0], 
                 'coef': float(match[1]), 
                 'std_err': float(match[2]), 
                 't_value': float(match[3]), 
                 'p_value': float(match[4])} for match in matches]

    def interpret_coefficient(self, coef: Dict[str, float], updated_feature_names: List[str]) -> str:
        """
        Interprets a single coefficient.
        
        :param coef: Dictionary containing coefficient data.
        :param updated_feature_names: List of updated feature names after preprocessing.
        :return: Interpretation string for the coefficient.
        """
        variable = coef['variable']  # Use the variable label directly

        coef_value = coef['coef']
        significance = coef['p_value'] < 0.05
        
        direction = "increases" if coef_value > 0 else "decreases"
        significance_txt = "statistically significant" if significance else "not statistically significant"
        
        return (f"For every unit increase in {variable}, the outcome {direction} by "
                f"{abs(coef_value):.2f} units, holding other factors constant. This effect is {significance_txt}.\n")



# Assumptions Interpreter
class AssumptionsInterpreter(BaseInterpreter):
    def interpret(self) -> str:
        try:
            self.interpretation += self.interpret_residuals() or ""
            self.interpretation += self.interpret_multicollinearity() or ""
            return self.interpretation
        except Exception as e:
            logging.error(f"General failure in interpretation: {str(e)}")
            return ""

    def interpret_residuals(self) -> Optional[str]:
        try:
            skew = self._extract_value(r"Skew:\s+([\d.-]+)", "skewness")
            kurtosis = self._extract_value(r"Kurtosis:\s+([\d.-]+)", "kurtosis")

            if skew is None or kurtosis is None:
                return None

            skewness_interpretation = self._interpret_skewness(skew)
            kurtosis_interpretation = self._interpret_kurtosis(kurtosis)

            return (f"{skewness_interpretation} The kurtosis of residuals indicates they are {kurtosis_interpretation} "
                    f"with a kurtosis value of {kurtosis}.\n")

        except Exception as e:
            logging.error(f"Unhandled exception while interpreting residuals: {str(e)}")
            return None

    def _extract_value(self, pattern: str, value_name: str) -> Optional[float]:
        try:
            value = self.parse(pattern)
            return float(value)
        except ValueError:
            logging.warning(f"Could not convert {value_name} to float. Check the output format.")
            return None
        except Exception as e:
            logging.error(f"Failed to extract {value_name}: {str(e)}")
            return None

    def _interpret_skewness(self, skew: float) -> str:
        if abs(skew) < 0.1:
            return "The residuals are approximately symmetrically distributed."
        elif skew > 0.1:
            return f"The residuals are right-skewed (positively skewed) with a skewness of {skew}."
        else:
            return f"The residuals are left-skewed (negatively skewed) with a skewness of {skew}."

    def _interpret_kurtosis(self, kurtosis: float) -> str:
        if abs(kurtosis - 3) < 0.5:
            return "approximately normal."
        elif kurtosis > 3.5:
            return "leptokurtic, indicating they have heavier tails and a sharper peak than a normal distribution."
        else:
            return "platykurtic, indicating they have thinner tails and a flatter peak than a normal distribution."

    def interpret_multicollinearity(self) -> Optional[str]:
        """
        Interprets the multicollinearity based on the condition number.
        
        :return: Interpretation string for multicollinearity or None if interpretation failed.
        """
        try:
            cond_number = self._extract_value(r"Cond. No.\s+([\d.-]+)", "condition number")
            
            if cond_number is None:
                return None

            if cond_number > 30:
                return (f"There might be a multicollinearity problem in the model, "
                        f"as indicated by the high condition number of {cond_number}.")
            else:
                return (f"There doesn't appear to be a multicollinearity problem in the model, "
                        f"as indicated by the condition number of {cond_number}.")

        except Exception as e:
            logging.error(f"Unhandled exception while interpreting multicollinearity: {str(e)}")
            return None


class Visualization:
    def __init__(self) -> None:
        if not os.path.exists('visuals'):
            os.makedirs('visuals')

    def scatter_plot(self, y: pd.Series, y_pred: pd.Series) -> None:
        """
        Creates a scatter plot of predicted vs actual values.

        :param y: Actual values.
        :param y_pred: Predicted values.
        """
        try:
            plt.scatter(y, y_pred, color='blue')
            plt.plot([y.min(), y.max()], [y.min(), y.max()], 'k--', lw=2)
            plt.xlabel('Actual Values')
            plt.ylabel('Predicted Values')
            plt.title('Scatter Plot of Predicted vs Actual Values')
            plt.savefig('visuals/scatter_plot.png')
            plt.show()
        except Exception as e:
            logging.error(f"An error occurred while generating scatter plot: {str(e)}")
            raise

    def pair_plot(self, data: pd.DataFrame) -> None:
        """
        Creates a pair plot of all variables in the data.

        :param data: Dataset containing the variables.
        """
        try:
            sns.pairplot(data)
            plt.savefig('visuals/pair_plot.png')
            plt.show()
        except Exception as e:
            logging.error(f"An error occurred while generating pair plot: {str(e)}")
            raise

    def residual_plot(self, y: pd.Series, y_pred: pd.Series) -> None:
        """
        Creates a residual plot.

        :param y: Actual values.
        :param y_pred: Predicted values.
        """
        try:
            residuals = y - y_pred
            plt.scatter(y_pred, residuals, color='red')
            plt.axhline(y=0, color='black', linestyle='--')
            plt.xlabel('Predicted Values')
            plt.ylabel('Residuals')
            plt.title('Residual Plot')
            plt.savefig('visuals/residual_plot.png')  # Save the plot
            plt.show()
        except Exception as e:
            logging.error(f"An error occurred while generating residual plot: {str(e)}")
            raise


class RegressionModel:
    """
    A class for running different types of regression models.
    """
    def __init__(self, model_type: str = 'linear') -> None:
        """
        Initializes the RegressionModel class with specified configurations.

        :param model_type: Type of regression model ('linear', 'ridge', or 'lasso').
        """
        self.model_type = model_type

    def create_model(self) -> Union[LinearRegression, Ridge, Lasso]:
        """
        Creates the regression model based on the specified model type.

        :return: Regression model object.
        """
        try:
            if self.model_type == 'linear':
                return LinearRegression()
            elif self.model_type == 'ridge':
                return Ridge()
            elif self.model_type == 'lasso':
                return Lasso()
            else:
                logging.error("Invalid model type selected.")
                raise ValueError("Invalid model type")
        except Exception as e:
            logging.error(f"An error occurred while creating the model: {str(e)}")
            raise

    def run(self, X_train: Any, X_test: Any, y_train: Any, y_test: Any, feature_names: List[str]) -> Tuple[Any, Any, dict]:
        """
        Main function to run the entire regression analysis process.

        :param X_train: Training set of independent variables.
        :param X_test: Testing set of independent variables.
        :param y_train: Training set of dependent variable.
        :param y_test: Testing set of dependent variable.
        :param feature_names: Names of the features.
        :return: Trained model, predicted values, and coefficients.
        """
        try:
            model = self.create_model()
            model.fit(X_train, y_train)
            y_pred = model.predict(X_test)
            
            # Check if the length of feature_names matches the number of features in X_train
            if len(feature_names) != X_train.shape[1]:
                raise ValueError("Length of feature_names does not match number of features in X_train")

            coefficients = dict(zip(feature_names, model.coef_))
            
            mse = mean_squared_error(y_test, y_pred)
            logging.info(f"Mean Squared Error: {mse}")
            
            return model, y_pred, coefficients
        except Exception as e:
            logging.error(f"An error occurred while running the model: {str(e)}")
            raise


if __name__ == "__main__":
    try:

        # 1. Load the data
        logging.info("Step 1: Loading data...")
        data_handler = DataHandler()
        file_path = ''  # replace with your data file path
        data = data_handler.load_data(file_path)
        
        # Assume 'y_column_name' is the name of your dependent variable column
        y = data.iloc[:, 0]  # First column as dependent variable
        X = data.iloc[:, 1:]  # All other columns as independent variables

        # 2. Preprocess the data
        logging.info("Step 2: Preprocessing data...")
        X_processed, y, updated_feature_names = data_handler.preprocess_data(X, y)

        # 3. Run regression
        logging.info("Step 3: Running regression...")
        analysis = StatisticalAnalysis(X_processed, y)
        model_stats = analysis.perform_analysis(updated_feature_names)
        
        # 4. Visualize results
        logging.info("Step 4: Generating visualizations...")
        visualizer = Visualization()
        y_pred = model_stats.predict(sm.add_constant(X_processed))  # Predict using the entire dataset
        visualizer.scatter_plot(y, y_pred)
        visualizer.residual_plot(y, y_pred)

        # 5. Generate statistical interpretation
        logging.info("Step 5: Generating statistical interpretation...")
        analysis = StatisticalAnalysis(X_processed, y)  # Use the entire processed dataset
        model_stats = analysis.perform_analysis(updated_feature_names)

        model_fit_interpreter = ModelFitInterpreter(model_stats.summary().as_text())
        #print(model_fit_interpreter.interpret())

        variable_interpreter = VariableInterpreter(model_stats.summary().as_text())
        #print(variable_interpreter.interpret(feature_names=updated_feature_names))

        # 6. Generate a report
        logging.info("Step 6: Generating report...")
        report = model_fit_interpreter.get_interpretation() + variable_interpreter.get_interpretation()
        ReportGenerator.save_to_word_doc(report)


    except Exception as e:
        logging.error(f"An error occurred: {str(e)}")
